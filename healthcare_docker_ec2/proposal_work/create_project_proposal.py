from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "/Users/abhinav/Documents/Codex/2026-05-01/please-guide-me-alternative-assessment-options/Cloud_Based_Healthcare_Clinical_Conversation_Analysis_Project_Proposal.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(10)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(2)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, header=True):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
            if header and row_idx == 0:
                set_cell_shading(cell, "EAF2F8")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.85)
section.bottom_margin = Inches(0.85)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(11)
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Normal"].paragraph_format.line_spacing = 1.08

for style_name, size, color in [
    ("Title", 22, RGBColor(31, 78, 121)),
    ("Heading 1", 15, RGBColor(31, 78, 121)),
    ("Heading 2", 12, RGBColor(47, 85, 151)),
]:
    style = styles[style_name]
    style.font.name = "Arial"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color

header = section.header
hp = header.paragraphs[0]
hp.text = "Practical Cloud Computing Project Proposal"
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in hp.runs:
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fp.text = "Cloud-Based Healthcare Clinical Conversation Analysis"
for run in fp.runs:
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("Cloud Deployment of an AI-Based Healthcare Clinical Conversation Analysis System")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Practical Cloud Computing Project Proposal")
run.font.name = "Arial"
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(90, 90, 90)

doc.add_paragraph()
meta = doc.add_table(rows=4, cols=2)
meta.columns[0].width = Inches(2.2)
meta.columns[1].width = Inches(4.9)
rows = [
    ("Assessment Option", "Practical Cloud Computing Project"),
    ("Project Area", "Cloud-based AI application deployment for healthcare conversation analysis"),
    ("Proposed Cloud Platform", "AWS EC2 with Docker deployment; optional AWS S3 and CloudWatch integration"),
    ("Core Technologies", "Python, Streamlit, Whisper, spaCy, VADER, scikit-learn, Sentence Transformers, Groq API, Docker"),
]
for i, (label, value) in enumerate(rows):
    set_cell_text(meta.cell(i, 0), label, bold=True)
    set_cell_text(meta.cell(i, 1), value)
    set_cell_shading(meta.cell(i, 0), "EAF2F8")
style_table(meta, header=False)

doc.add_heading("1. Project Objective", level=1)
doc.add_paragraph(
    "The objective of this project is to deploy an AI-powered healthcare conversation analysis application in a cloud environment. "
    "The system will allow users to upload clinical conversation files, either text or audio, and receive structured insights such as transcription, sentiment analysis, named entities, speaker segmentation, automated summaries, and question-answering results."
)
doc.add_paragraph(
    "The expected outcome is a working cloud-hosted Streamlit application that demonstrates how cloud computing can make advanced NLP and machine learning services accessible through a browser-based interface."
)

doc.add_heading("2. Relevance to Cloud Computing", level=1)
doc.add_paragraph(
    "This project aligns with cloud computing principles because the application is designed to be deployed, accessed, and tested through a remote cloud environment rather than only running on a local machine. "
    "The cloud deployment supports availability, remote access, compute provisioning, API integration, and future scalability."
)
for item in [
    "Cloud compute will be used to host and run the Streamlit application.",
    "Docker will package the application and its dependencies for repeatable deployment.",
    "Environment variables will be used to manage API credentials such as the Groq API key.",
    "Cloud networking will expose the web application through a controlled public endpoint.",
    "Optional cloud storage can be used for sample input files, processed transcripts, or project artifacts.",
    "Optional monitoring can track application health, logs, and runtime behavior.",
]:
    add_bullet(doc, item)

doc.add_heading("3. Problem Statement", level=1)
doc.add_paragraph(
    "Clinical conversations contain important information about symptoms, patient concerns, diagnosis, treatment decisions, and follow-up actions. "
    "However, these conversations are often documented manually, which can be time-consuming and may cause useful insights to be missed. "
    "Traditional patient feedback also depends heavily on surveys, which are subjective and usually collected after the visit."
)
doc.add_paragraph(
    "This project addresses that problem by creating a cloud-accessible system that can analyze clinical conversations automatically and produce useful healthcare intelligence from conversation data."
)

doc.add_heading("4. Proposed Solution", level=1)
doc.add_paragraph(
    "The proposed solution is a cloud-deployed Streamlit application for clinical conversation analysis. "
    "Users upload an audio or text file, and the application processes it through a multi-stage AI pipeline."
)
for item in [
    "Whisper converts uploaded audio conversations into text.",
    "VADER calculates sentiment scores to estimate the emotional tone of the conversation.",
    "spaCy extracts entities such as people, organizations, dates, and locations.",
    "A TF-IDF and Logistic Regression model classifies likely speaker roles such as Doctor and Patient.",
    "Sentence Transformers support semantic retrieval for transcript-based question answering.",
    "The Groq LLM API generates structured summaries and answers user questions based on retrieved context.",
]:
    add_bullet(doc, item)

doc.add_heading("5. Methodology", level=1)
for item in [
    "Finalize the existing Streamlit application and align labels, headings, and sample data with the healthcare use case.",
    "Prepare synthetic or sample clinical conversation files for testing. Real patient data will not be used.",
    "Create a requirements file and Dockerfile so the application can be deployed consistently.",
    "Configure the Groq API key as a cloud environment variable instead of hardcoding credentials.",
    "Deploy the containerized application on AWS EC2.",
    "Open only the required network port for the Streamlit web interface and test browser access.",
    "Run sample uploads and capture screenshots of transcription, sentiment, entities, speaker segmentation, summary generation, and chatbot responses.",
    "Document results, limitations, challenges, and future improvements in the final report.",
]:
    add_number(doc, item)

doc.add_heading("6. System Architecture", level=1)
arch = doc.add_paragraph()
arch.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = arch.add_run(
    "User Browser -> Cloud-hosted Streamlit App -> AI/NLP Processing Pipeline -> Groq LLM API -> Results Dashboard"
)
r.bold = True
r.font.name = "Arial"
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(31, 78, 121)
doc.add_paragraph(
    "The architecture separates the user interface, application logic, machine learning models, and external LLM service. "
    "The cloud server hosts the Streamlit application and model dependencies, while Groq provides LLM-based summarization and question answering through an API call."
)

doc.add_heading("7. Resources and Tools Required", level=1)
tools = doc.add_table(rows=1, cols=2)
set_cell_text(tools.cell(0, 0), "Category", bold=True)
set_cell_text(tools.cell(0, 1), "Tools / Resources", bold=True)
tool_rows = [
    ("Programming and Interface", "Python, Streamlit"),
    ("Speech-to-Text", "Whisper base model"),
    ("NLP and Sentiment", "spaCy, VADER Sentiment"),
    ("Machine Learning", "scikit-learn, TF-IDF, Logistic Regression"),
    ("Semantic Retrieval", "Sentence Transformers all-MiniLM-L6-v2"),
    ("LLM Integration", "Groq API using llama-3.3-70b-versatile"),
    ("Cloud Deployment", "AWS EC2, Docker, security groups, environment variables"),
    ("Testing and Documentation", "Sample clinical conversations, screenshots, final report, demo"),
]
for category, value in tool_rows:
    cells = tools.add_row().cells
    set_cell_text(cells[0], category, bold=True)
    set_cell_text(cells[1], value)
style_table(tools)

doc.add_heading("8. Tentative Timeline", level=1)
timeline = doc.add_table(rows=1, cols=3)
for i, text in enumerate(["Week", "Milestone", "Deliverable"]):
    set_cell_text(timeline.cell(0, i), text, bold=True)
timeline_rows = [
    ("Week 1", "Finalize project scope and prepare proposal", "Approved project proposal"),
    ("Week 2", "Clean application code and prepare test files", "Local working application with sample inputs"),
    ("Week 3", "Create Dockerfile, requirements file, and environment setup", "Container-ready application"),
    ("Week 4", "Deploy the application on AWS EC2", "Cloud-hosted Streamlit application"),
    ("Week 5", "Test all features and collect evidence", "Screenshots, test results, deployment notes"),
    ("Week 6", "Prepare final report and demo", "Final submission and practical demonstration"),
]
for row in timeline_rows:
    cells = timeline.add_row().cells
    for i, text in enumerate(row):
        set_cell_text(cells[i], text, bold=(i == 0))
style_table(timeline)

doc.add_heading("9. Expected Outcomes", level=1)
for item in [
    "A deployed cloud-based healthcare conversation analysis application.",
    "Successful upload and analysis of text and audio conversation files.",
    "Speech-to-text transcription for audio inputs.",
    "Sentiment analysis output showing the overall emotional tone of the conversation.",
    "Entity extraction for people, organizations, dates, and locations.",
    "Doctor-patient speaker segmentation using a lightweight machine learning model.",
    "LLM-generated structured summaries and action items.",
    "RAG-based question answering over the analyzed conversation.",
    "A final report with architecture, deployment steps, screenshots, results, limitations, and future work.",
]:
    add_bullet(doc, item)

doc.add_heading("10. Evaluation Criteria Alignment", level=1)
criteria = doc.add_table(rows=1, cols=2)
set_cell_text(criteria.cell(0, 0), "Criterion", bold=True)
set_cell_text(criteria.cell(0, 1), "How the Project Addresses It", bold=True)
criteria_rows = [
    ("Practical application of cloud concepts", "The application is packaged and deployed on a cloud compute instance with remote browser access."),
    ("Innovation and originality", "The project combines speech-to-text, NLP, sentiment analysis, speaker classification, RAG, and LLM summarization for healthcare conversation analysis."),
    ("Technical proficiency", "The system uses multiple AI/ML libraries, API integration, containerization, and cloud deployment practices."),
    ("Project management and execution", "The timeline divides the work into proposal, development, deployment, testing, and final documentation phases."),
    ("Impact and usefulness", "The system can help extract insights from clinical conversations and reduce dependence on manual review or subjective surveys."),
]
for criterion, alignment in criteria_rows:
    cells = criteria.add_row().cells
    set_cell_text(cells[0], criterion, bold=True)
    set_cell_text(cells[1], alignment)
style_table(criteria)

doc.add_heading("11. Limitations and Ethical Considerations", level=1)
for item in [
    "The project will use synthetic or sample clinical conversations only. Real patient data will not be used.",
    "The application is a prototype and should not be treated as a medical decision-making system.",
    "Speaker classification is trained on a small sample and may need more data for production-level accuracy.",
    "LLM-generated summaries should be reviewed by a human before being used in any clinical context.",
    "Future work could include stronger authentication, encrypted storage, HIPAA-aware deployment design, model monitoring, and improved clinical entity extraction.",
]:
    add_bullet(doc, item)

doc.add_heading("12. Conclusion", level=1)
doc.add_paragraph(
    "This project demonstrates how cloud computing can support practical AI applications in healthcare. "
    "By deploying the clinical conversation analysis system in the cloud, the project shows how speech-to-text, NLP, machine learning, and LLM services can be combined into an accessible web application. "
    "The final submission will include the deployed application, testing evidence, architecture explanation, and a practical demonstration of the system."
)

doc.save(OUTPUT)
print(OUTPUT)
